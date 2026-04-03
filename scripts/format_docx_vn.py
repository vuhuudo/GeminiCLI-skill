import os
import sys
import re
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

SIGNATURE_KEYWORDS = ["người lập", "giám đốc", "phê duyệt", "trưởng phòng", "kế toán"]


def _set_run_font(run, size_pt, bold=None, italic=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_docx(input_path, output_path):
    if not os.path.exists(input_path):
        print(f"Error: File '{input_path}' does not exist.")
        sys.exit(1)

    doc = Document(input_path)

    # --- SMART ANALYSIS ---
    paragraphs = list(doc.paragraphs)
    total_chars = sum(len(p.text) for p in paragraphs)
    total_paragraphs = len(paragraphs)

    # Determine if we need to be aggressive to save space
    is_long = total_chars > 3200 or total_paragraphs > 40

    font_size_pt = 13 if is_long else 14
    line_spacing = 1.0 if is_long else 1.15
    space_after = Pt(2) if is_long else Pt(6)

    # 1. Set Margins (Nghị định 30)
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)

    # 2. Set Default Styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_size_pt)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.first_line_indent = Mm(10)

    # 3. Process Paragraphs
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()

        # Remove trailing empty paragraphs (last 10% of document)
        if not text and i > len(paragraphs) * 0.9:
            p_element = paragraph._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            continue

        if not text:
            continue

        # Reset to Normal style
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            _set_run_font(run, font_size_pt)

        # --- SPECIAL FORMATTING RULES ---

        # Quốc hiệu
        if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text.upper():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.space_after = 0
            for run in paragraph.runs:
                _set_run_font(run, 12, bold=True)

        # Tiêu ngữ
        elif re.search(r"ĐỘC LẬP[\s-]+TỰ DO[\s-]+HẠNH PHÚC", text.upper()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.space_after = Pt(10)
            for run in paragraph.runs:
                _set_run_font(run, 13, bold=True)

        # Đề mục (I., II., 1., 2., a) ...) - No Indent, Bold
        elif re.match(r'^([IVXLCDM]+\.|[0-9]+\.|[a-z]\))\s+', text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.space_before = Pt(10)
            for run in paragraph.runs:
                _set_run_font(run, font_size_pt + 1, bold=True)

        # Địa danh, ngày tháng
        elif re.search(r",\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+", text, re.IGNORECASE):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = 0
            for run in paragraph.runs:
                _set_run_font(run, font_size_pt, italic=True)

        # Tiêu đề văn bản (ALL CAPS, short)
        elif text.isupper() and len(text) < 100:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                _set_run_font(run, font_size_pt + 1, bold=True)

        # Chữ ký (Signature block detection — checked independently, may override above)
        if any(k in text.lower() for k in SIGNATURE_KEYWORDS) and len(text) < 40:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.space_before = Pt(10) if is_long else Pt(20)
            for run in paragraph.runs:
                _set_run_font(run, font_size_pt, bold=True)

    # 4. Process Tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = 0
                    paragraph.paragraph_format.space_after = 0
                    if any(k in paragraph.text.lower() for k in ["giám đốc", "người lập"]):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        _set_run_font(run, 12)

    try:
        doc.save(output_path)
    except OSError as e:
        print(f"Error: Could not save file '{output_path}': {e}")
        sys.exit(1)

    print(f"Formatted successfully: {output_path} (Long mode: {is_long})")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python format_docx_vn.py <input_path> <output_path>")
        sys.exit(1)
    format_docx(sys.argv[1], sys.argv[2])
